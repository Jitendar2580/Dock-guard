import os
import tempfile
from flask import Flask, request, jsonify, render_template
from dotenv import load_dotenv
from docx import Document as DocxDocument
from pdf2image import convert_from_bytes
import pytesseract

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import PyPDFLoader
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from langchain_core.prompts import PromptTemplate
from langchain_core.documents import Document
from langchain.memory import ConversationBufferMemory


# ---------------------
# Flask setup
# ---------------------
app = Flask(__name__)
load_dotenv()

# OCR setup (skip on Linux if Render has tesseract installed)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Embeddings and model
embedding_model = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')
memory = ConversationBufferMemory()
client = ChatGroq(model="llama-3.3-70b-versatile", api_key=os.getenv("API_KEY"))

# Global retriever
retriever = None


# ---------------------
# Helpers
# ---------------------
def extract_text_from_scanned_pdf(file_bytes):
    text = ""
    images = convert_from_bytes(file_bytes)
    for img in images:
        text += pytesseract.image_to_string(img)
    return text


def process_document(file):
    file_bytes = file.read()
    file_name = file.filename.lower()

    if file_name.endswith(".txt"):
        content = file_bytes.decode("utf-8")
        docs = [Document(page_content=content)]
    elif file_name.endswith(".pdf"):
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            loader = PyPDFLoader(tmp_path)
            docs = loader.load()
        except Exception:
            text = extract_text_from_scanned_pdf(file_bytes)
            docs = [Document(page_content=text)]
    elif file_name.endswith(".docx"):
        docx = DocxDocument(file)
        full_text = "\n".join([p.text for p in docx.paragraphs])
        docs = [Document(page_content=full_text)]
    else:
        docs = []

    return docs


# ---------------------
# Routes
# ---------------------
@app.route("/")
def home():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    global retriever
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    docs = process_document(file)

    if not docs:
        return jsonify({"error": "Unsupported or empty file"}), 400

    # Chunking + FAISS
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(docs)
    db = FAISS.from_documents(chunks, embedding_model)
    retriever = db.as_retriever()

    return jsonify({"message": "Document uploaded and processed"})


@app.route("/chat", methods=["POST"])
def chat():
    global retriever
    if not retriever:
        return jsonify({"error": "No document uploaded"}), 400

    user_input = request.json.get("question")
    if not user_input:
        return jsonify({"error": "No question provided"}), 400

    # Prompt
    # prompt_template = """
    #     You are DocGuard AI - Smart Document Assistant for Risk Detection & Compliance Assurance.

    #     Here's the relevant context from the document:
    #     ---------------------
    #     {context}
    #     ---------------------

    #     When responding:
    #     - Highlight risky clauses with severity ratings
    #     - Point out compliance gaps
    #     - Suggest improvements
    #     - Be professional and concise
    #     - If you don’t know, say so
    #     User: {input}
    #     DocGuard AI:
    # """
    
    prompt_template ="""
        You are an expert document analyst. I will provide you with a document (any format: policy, report, training pack, compliance doc, or guideline).  
        
        Here's the relevant context from the document:
        ---------------------
        {context}
        ---------------------

        Your task is to carefully read and analyze the document, then provide the following structured output:

        1. **Summary**  
        - Concise overview of the document in 5–7 bullet points.  

        2. **Key Details & Insights**  
        - Highlight important facts, rules, requirements, or findings.  
        - Extract any numbers, obligations, or deadlines if present.  

        3. **Intention of the Document**  
        - Clearly explain the purpose behind this document.  
        - Who is it for? Why does it exist? What problem does it solve?  

        4. **Solutions / Guidance**  
        - Based on the document, suggest how it should be applied in practice.  
        - If it’s a compliance or regulatory doc, explain what must be done to meet requirements.  
        - If it’s a policy, describe how to implement it.  
        - If it’s a learning/training doc, explain how learners or assessors should use it.  

        5. **Plain English Explanation**  
        - Simplify the document’s meaning so anyone without technical knowledge can understand it.  

        6. **Actionable Points**  
        - Provide a checklist or next steps that the reader should follow.  

        Make sure the response is well-structured, professional, and easy to understand.
        
        
        
         When responding:
        - Highlight risky clauses with severity ratings
        - Point out compliance gaps
        - Suggest improvements
        - Be professional and concise
        - If you don’t know, say so
        User: {input}
        DocGuard AI:
    """
    prompt = PromptTemplate(input_variables=["context", "input"], template=prompt_template)
    document_chain = create_stuff_documents_chain(client, prompt)
    retrieval_chain = create_retrieval_chain(retriever, document_chain)

    response = retrieval_chain.invoke({"input": user_input})
    return jsonify({"answer": response["answer"]})


# ---------------------
# Run
# ---------------------
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
