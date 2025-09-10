
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import TextLoader, PyPDFLoader
from sentence_transformers import SentenceTransformer
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from langchain_core.prompts import PromptTemplate
from langchain_core.documents import Document
from docx import Document as DocxDocument
from pdf2image import convert_from_bytes
import pytesseract
import tempfile
import os
from dotenv import load_dotenv
from langchain.memory import ConversationBufferMemory


load_dotenv() 

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# Embedding and LLM setup
embedding_model = HuggingFaceEmbeddings(model_name='sentence-transformers/all-MiniLM-L6-v2')

memory=ConversationBufferMemory()
client = ChatGroq(model="llama-3.3-70b-versatile", api_key=os.getenv("API_KEY"))

# Prompt setup for chatbot interaction
prompt_template = """
                  You are DocGuard AI - Smart Document Assistant for Risk Detection & Compliance Assurance.
                  
                  You're having a conversation with a user about their document. Here's the relevant context from their document:
                  ---------------------
                  {context}
                  ---------------------
                  
                  When responding:
                  1. First understand the user's question/request
                  2. Refer to the document context when relevant
                  3. For risk/compliance questions:
                     - Highlight any risky clauses with severity ratings
                     - Point out compliance gaps if found
                     - Suggest improvements when appropriate
                  4. Keep responses conversational but professional
                  5. If you don't know, say so - don't make up information
                  
                  Current conversation:
                  User: {input}
                  DocGuard AI: """

prompt = PromptTemplate(input_variables=["context", "input"], template=prompt_template)
document_chain = create_stuff_documents_chain(client, prompt)

# OCR for scanned PDFs
def extract_text_from_scanned_pdf(file_bytes):
	text = ""
	images = convert_from_bytes(file_bytes)
	for img in images:
		text += pytesseract.image_to_string(img)
	return text

# Split and embed documents
def process_document(file):
	file_bytes = file.read()
	file_name = file.name.lower()

	if file_name.endswith(".txt"):
		content = file_bytes.decode("utf-8")
	elif file_name.endswith(".pdf"):
		try:
			loader = PyPDFLoader(file.name)
			docs = loader.load()
			return docs
		except Exception:
			# Try OCR
			text = extract_text_from_scanned_pdf(file_bytes)
			return [Document(page_content=text)]
	elif file_name.endswith(".docx"):
		docx = DocxDocument(file)
		full_text = "\n".join([p.text for p in docx.paragraphs])
		return [Document(page_content=full_text)]
	else:
		return []

	return [Document(page_content=content)]

if "messages" not in st.session_state:
	st.session_state.messages = []

st.set_page_config(page_title="DocGuard AI Chat", layout="wide")
st.title("💬 DocGuard AI - Document Chat Assistant")

with st.sidebar:
	st.header("📄 Upload Document")
	uploaded_file = st.file_uploader("Upload your document (TXT, PDF, DOCX)", type=["txt", "pdf", "docx"], key="file_uploader")

	if uploaded_file and "processed_doc" not in st.session_state:
		with st.spinner("Processing document..."):
			docs = process_document(uploaded_file)
			if docs:
				# Chunking
				splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
				chunks = splitter.split_documents(docs)

				# Vector DB
				db = FAISS.from_documents(chunks, embedding_model)
				st.session_state.retriever = db.as_retriever()
				st.session_state.processed_doc = True
				st.success("Document processed and ready for questions!")
			else:
				st.error("Failed to process document")

for message in st.session_state.messages:
	with st.chat_message(message["role"]):
		st.markdown(message["content"])

if prompt := st.chat_input("Ask me about your document..."):
	if "processed_doc" not in st.session_state:
		st.warning("Please upload a document first")
		st.stop()

	st.session_state.messages.append({"role": "user", "content": prompt})

	with st.chat_message("user"):
		st.markdown(prompt)

	with st.chat_message("assistant"):
		with st.spinner("Analyzing..."):
			retrieval_chain = create_retrieval_chain(st.session_state.retriever, document_chain)
			response = retrieval_chain.invoke({"input": prompt})
			st.markdown(response["answer"])
	st.session_state.messages.append({"role": "assistant", "content": response["answer"]})
